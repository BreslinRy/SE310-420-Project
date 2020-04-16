from docx import Document
import csv
from docx2pdf import convert


#I should comment this code
def genReceipt(VoterData):
    voterID = 0
    receipt = Document()
    header = receipt.add_heading('Voters Receipt')
    header.bold = True
    header.alignment = 1
    with open(VoterData) as csvFile:
        lineCount = 1
        csvReader = csv.reader(csvFile, delimiter=',')
        for row in csvReader:
            if lineCount == 1:
                voterID = row[0]
                intro = receipt.add_paragraph('This receipt corresponds to voter ID: ')
                intro.add_run(voterID)
                intro.add_run('.')
            if lineCount == int(row[1]):
                if row[2] == 'yes' or row[2] == 'no':
                    vote = receipt.add_paragraph('You have voted ')
                    vote.add_run(row[2])
                    vote.add_run(' for proposition ')
                    vote.add_run(row[3])
                    vote.add_run('.')
                else:
                    vote = receipt.add_paragraph('You have voted for ')
                    vote.add_run(row[2])
                    vote.add_run(' for the position of ')
                    vote.add_run(row[3])
                    vote.add_run('.')
                lineCount = lineCount + 1
            else:
                receipt = Document()
                header = receipt.add_heading('INVALID BALLOT')
                header.bold = True
                header.alignment = 1
                error = receipt.add_paragraph('An error has been encountered with your ballot please contact your polling station to resolve the issue for Voter ID: ')
                error.add_run(voterID)
                error.add_run('.')
                receipt.save('VotingReceipt.docx')
                convert("VotingReceipt.docx")
                return "Failed due to inconsistent ballot data"
    receipt.add_paragraph('Thank you for voting in this years election. Your votes have been recorded and this document acts as proof. Please save and print this file to maintain a paper trail of this interaction.')
    receipt.save('VotingReceipt.docx')
    convert("VotingReceipt.docx")
    return "Success"


print(genReceipt('test1.csv'))
input("press enter to continue")
print(genReceipt('test2.csv'))